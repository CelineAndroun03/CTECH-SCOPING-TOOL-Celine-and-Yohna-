
"""
Exploratory Data Analysis (EDA) – Dual Targets: Eng. AH & Lab. AH
Author: Nuno Monteiro and Celine Androun

Purpose:
- Perform a robust EDA for two target variables: "Eng. AH" and "Lab. AH".
- Save ALL outputs (printed text, key tables, all plots/figures) into a single Word (.docx) report.

Usage:
    python eda_report.py --input "1. Data.xlsx" --output "EDA_Report.docx"

Notes:
- All figures are saved to ./reports/figures and embedded into the Word report with captions.
"""

# ==================================================
# 0. Imports & Global Configuration
# ==================================================
import argparse
import io
import os
from pathlib import Path
from datetime import datetime
import warnings
import numpy as np
import pandas as pd
import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
import seaborn as sns
from scipy.stats import median_abs_deviation, zscore
from scipy import stats as sstats
from sklearn.metrics.pairwise import pairwise_distances
from sklearn.feature_selection import mutual_info_regression, mutual_info_classif
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_BREAK
from contextlib import redirect_stdout

# Global display and warnings
warnings.filterwarnings("ignore")
pd.set_option("display.max_columns", None)
pd.set_option("display.max_rows", None)
sns.set_theme(style="whitegrid")

# Reproducibility (where applicable)
RANDOM_STATE = 42
np.random.seed(RANDOM_STATE)

# ==================================================
# 1. Utilities: logging, files, figures, report
# ==================================================
class ReportArtifacts:
    """
    Container to track outputs during EDA to later assemble the Word report.
    """
    def __init__(self, base_dir: Path):
        self.base_dir = base_dir
        self.fig_dir = base_dir / "figures"
        self.fig_dir.mkdir(parents=True, exist_ok=True)
        self.images = []  # list of dicts: {"path": Path, "caption": str}
        self.tables = {}  # dict of {name: DataFrame}
        self.stdout_text = ""  # captured printed text

    def save_fig(self, plt_obj, caption: str, fname: str = None, dpi: int = 150):
        """
        Save a matplotlib figure to the figures directory and register it with caption.
        """
        if fname is None:
            # Make a file-name friendly version of caption
            safe = "".join(c for c in caption if c.isalnum() or c in ("-", "_", " ")).rstrip()
            safe = safe.replace(" ", "_")[:80]
            fname = f"{safe}.png"
        out_path = self.fig_dir / fname
        plt_obj.savefig(out_path, dpi=dpi, bbox_inches="tight")
        plt_obj.close()
        self.images.append({"path": out_path, "caption": caption})

    def add_table(self, name: str, df: pd.DataFrame):
        self.tables[name] = df.copy()


def capture_stdout(func):
    """
    Decorator to capture everything printed to stdout within the wrapped function.
    """
    def wrapper(*args, **kwargs):
        artifacts: ReportArtifacts = kwargs.get("artifacts", None)
        if artifacts is None and len(args) > 0:
            # By convention, assume artifacts is the last arg if provided positionally
            for a in args:
                if isinstance(a, ReportArtifacts):
                    artifacts = a
                    break

        buf = io.StringIO()
        with pd.option_context("display.width", 120):
            with redirect_stdout(buf):
                return_val = func(*args, **kwargs)
        if artifacts:
            artifacts.stdout_text += buf.getvalue() + "\n"
        return return_val
    return wrapper

def write_word_report(artifacts: ReportArtifacts, output_docx: Path, title: str, meta: dict):
    """
    Compose the Word report with:
    - Title page
    - Summary metadata
    - Selected tables (as Word tables)
    - All saved figures (with captions)
    - Appendix with all printed text captured during EDA
    """
    doc = Document()

    # ---- Title ----
    doc.add_heading(title, level=0)
    p = doc.add_paragraph()
    p.add_run(f"Generated on: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n")
    for k, v in meta.items():
        p.add_run(f"{k}: {v}\n")

    doc.add_page_break()

    # ---- Tables ----
    if artifacts.tables:
        doc.add_heading("Key Tables", level=1)
        for name, df in artifacts.tables.items():
            doc.add_heading(name, level=2)
            MAX_ROWS = 500
            truncated = False
            df_show = df.copy()
            if len(df_show) > MAX_ROWS:
                df_show = df_show.head(MAX_ROWS)
                truncated = True

            # Create Word table
            n_rows, n_cols = df_show.shape
            table = doc.add_table(rows=1, cols=n_cols)
            table.style = "Light List Accent 1"
            hdr_cells = table.rows[0].cells
            for j, col in enumerate(df_show.columns):
                hdr_cells[j].text = str(col)

            for _, row in df_show.iterrows():
                row_cells = table.add_row().cells
                for j, col in enumerate(df_show.columns):
                    # Convert values to string, formatted where possible
                    val = row[col]
                    if isinstance(val, float):
                        row_cells[j].text = f"{val:.6g}"
                    else:
                        row_cells[j].text = str(val)

            if truncated:
                doc.add_paragraph("Note: Table truncated for display (showing first 500 rows).")

            doc.add_paragraph()

        doc.add_page_break()

    # ---- Figures ----
    if artifacts.images:
        doc.add_heading("Figures & Visualizations", level=1)
        for item in artifacts.images:
            doc.add_paragraph(item["caption"]).runs[0].bold = True
            doc.add_picture(str(item["path"]), width=Inches(6.5))
            doc.add_paragraph()  # spacer
        doc.add_page_break()

    # ---- Appendix: Raw Printed Output ----
    if artifacts.stdout_text.strip():
        doc.add_heading("Appendix – Full Printed Output", level=1)
        para = doc.add_paragraph()
        run = para.add_run(artifacts.stdout_text)
        font = run.font
        font.name = "Courier New"
        font.size = Pt(9)

    # Save the document
    output_docx.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_docx))


# ==================================================
# 2. Data helpers
# ==================================================
def read_data(input_path: Path) -> pd.DataFrame:
    """
    Read Excel data with strong typing for target columns and safe coercion.
    """
    if not input_path.exists():
        raise FileNotFoundError(f"Input file not found: {input_path}")

    df = pd.read_excel(input_path, engine="openpyxl")

    # Drop non-target SH columns if present, ignore otherwise
    df = df.drop(columns=["Eng. SH", "Lab. SH"], errors="ignore")

    # Safe numeric coercion for known numeric columns
    if "total_test_count" in df.columns:
        df["total_test_count"] = pd.to_numeric(df["total_test_count"], errors="coerce")

    for col in ["Eng. AH", "Lab. AH"]:
        if col in df.columns:
            df[col] = pd.to_numeric(df[col], errors="coerce")
        else:
            # Keep running but warn
            print(f"[WARN] Missing required target column: {col}")

    return df


def iqr_outlier_stats(s: pd.Series, k: float = 1.5):
    """
    Compute IQR-based outlier statistics for a numeric series.
    """
    s = s.dropna()
    if s.empty:
        return {"q1": np.nan, "q3": np.nan, "iqr": np.nan, "lcl": np.nan, "ucl": np.nan,
                "n": 0, "n_outliers": 0, "pct_outliers": 0.0}
    q1, q3 = s.quantile([0.25, 0.75])
    iqr = q3 - q1
    low = q1 - k * iqr
    high = q3 + k * iqr
    outliers = s[(s < low) | (s > high)]
    return {
        "q1": q1, "q3": q3, "iqr": iqr,
        "lcl": low, "ucl": high,
        "n": s.size, "n_outliers": outliers.size,
        "pct_outliers": 100 * outliers.size / s.size if s.size > 0 else 0.0
    }


def z_outliers_idx(s: pd.Series, threshold: float = 3.0):
    """
    Get index of Z-score outliers.
    """
    x = s.dropna()
    if x.empty:
        return x.index[:0]
    z = zscore(x)
    return x.index[np.abs(z) > threshold]


def robust_z_outliers_idx(s: pd.Series, threshold: float = 3.5):
    """
    Get index of robust Z-score outliers using MAD.
    """
    x = s.dropna()
    if x.empty:
        return x.index[:0]
    med = np.median(x)
    mad = median_abs_deviation(x, scale='normal')
    denom = mad if mad != 0 else 1.0
    rz = 0.6745 * (x - med) / denom
    return x.index[np.abs(rz) > threshold]


def iqr_fences(s: pd.Series, k: float = 1.5):
    """
    Compute lower and upper IQR fences for a numeric series.
    """
    s = s.dropna()
    if s.empty:
        return np.nan, np.nan
    q1, q3 = s.quantile([0.25, 0.75])
    iqr = q3 - q1
    return q1 - k * iqr, q3 + k * iqr


def safe_log(x: pd.Series, min_positive: float = 1e-6) -> pd.Series:
    """
    Safe natural logarithm for nonpositive values by clipping to a tiny positive constant.
    """
    return np.log(x.clip(lower=min_positive))


def cramers_v(a: pd.Series, b: pd.Series) -> float:
    """
    Compute bias-corrected Cramér’s V for two categorical variables.
    """
    tbl = pd.crosstab(a, b)
    if tbl.size == 0:
        return np.nan
    chi2 = sstats.chi2_contingency(tbl, correction=False)[0]
    n = tbl.values.sum()
    r, k = tbl.shape
    if n == 0:
        return np.nan
    phi2 = chi2 / n
    phi2corr = max(0, phi2 - (k - 1) * (r - 1) / (n - 1)) if n > 1 else 0
    rcorr = r - (r - 1) ** 2 / (n - 1) if n > 1 else r
    kcorr = k - (k - 1) ** 2 / (n - 1) if n > 1 else k
    denom = max(1e-9, min((kcorr - 1), (rcorr - 1)))
    return np.sqrt(phi2corr / denom) if denom > 0 else 0.0


# ==================================================
# 3. EDA Steps (each step logs to stdout and saves figs/tables)
# ==================================================
@capture_stdout
def step_preliminary_info(df: pd.DataFrame, artifacts: ReportArtifacts):
    print("##########################################")
    print("##### Dataset Preliminary Information #####")
    print("##########################################\n")

    # Head
    print("First 5 dataset rows:")
    print(df.head(5), "\n")

    # Info
    print("Info Function:")
    buf = io.StringIO()
    df.info(buf=buf)
    print(buf.getvalue())

    # Describe
    # Only describe numeric columns for readability
    print("Describe (numeric columns):")
    print(df.describe(include=[np.number]), "\n")

    # Missing values table
    print("Missing values:")
    missing_counts = df.isna().sum()
    missing_pct = df.isna().mean() * 100
    missing_table = (
        pd.DataFrame({
            "Column": df.columns,
            "Number of Missing Values": missing_counts.values,
            "% of Missing Values": missing_pct.values
        }).sort_values("% of Missing Values", ascending=False).reset_index(drop=True)
    )
    print(missing_table, "\n")
    artifacts.add_table("Missing Values Summary", missing_table)

    # Unique values
    print("Number of unique values (per column):")
    print(df.nunique(), "\n")

    # Duplicates
    dup_count = df.duplicated().sum()
    print(f"Duplicated Rows: {dup_count}")
    if dup_count > 0:
        print("Dropping duplicated rows (keep='first').")
        df.drop_duplicates(keep="first", inplace=True)

    # Capture library versions
    print("\nLibrary Versions:")
    print(f"pandas: {pd.__version__}")
    print(f"numpy:  {np.__version__}")
    import seaborn as _sns; print(f"seaborn: {_sns.__version__}")
    import matplotlib as _mpl; print(f"matplotlib: {_mpl.__version__}")
    import scipy as _sc; print(f"scipy: {_sc.__version__}")
    import sklearn as _sk; print(f"scikit-learn: {_sk.__version__}")


def plot_hist_with_kde(series: pd.Series, title: str, xlabel: str, artifacts: ReportArtifacts, bins: int = 30, color: str = "#4C78A8"):
    """
    Utility to plot KDE-backed histogram and save.
    """
    fig, ax = plt.subplots(figsize=(10, 5))
    sns.histplot(series.dropna(), kde=True, bins=bins, ax=ax, color=color)
    ax.set_title(title)
    ax.set_xlabel(xlabel)
    ax.set_ylabel("Count")
    artifacts.save_fig(plt, caption=title)


@capture_stdout
def step_univariate(df: pd.DataFrame, artifacts: ReportArtifacts):
    print("\n##########################################")
    print("########## Univariate Analysis ###########")
    print("##########################################\n")

    targets = [c for c in ["Eng. AH", "Lab. AH"] if c in df.columns]
    if "Eng. AH" in targets:
        plot_hist_with_kde(df["Eng. AH"], "Distribution of Engineering Hours (Eng. AH)", "Eng. AH", artifacts)
    if "Lab. AH" in targets:
        plot_hist_with_kde(df["Lab. AH"], "Distribution of Laboratory Hours (Lab. AH)", "Lab. AH", artifacts, color="#74c476")

    # Outlier summaries
    for col in targets:
        stats_iqr = iqr_outlier_stats(df[col], k=1.5)
        print(f"\n[{col}] IQR-based outlier summary:")
        for k, v in stats_iqr.items():
            if isinstance(v, (int, float, np.floating)):
                print(f"{k:>12}: {v:.4f}")
            else:
                print(f"{k:>12}: {v}")

    # Boxplots
    if "Eng. AH" in df.columns and "Lab. AH" in df.columns:
        fig, axes = plt.subplots(1, 2, figsize=(12, 4))
        sns.boxplot(y=df["Eng. AH"], ax=axes[0], color="#6baed6")
        axes[0].set_title("Boxplot – Eng. AH")
        sns.boxplot(y=df["Lab. AH"], ax=axes[1], color="#74c476")
        axes[1].set_title("Boxplot – Lab. AH")
        plt.tight_layout()
        artifacts.save_fig(plt, caption="Boxplots – Eng. AH and Lab. AH")

    # Log-transformed distributions
    if "Eng. AH" in df.columns:
        fig, axes = plt.subplots(1, 2, figsize=(12, 4))
        sns.histplot(np.log(df["Eng. AH"].loc[df["Eng. AH"] > 0]), kde=True, bins=30, ax=axes[0])
        axes[0].set_title("Log Distribution – Eng. AH (x>0)")
        axes[0].set_xlabel("log(Eng. AH)")

        if "Lab. AH" in df.columns:
            sns.histplot(np.log1p(df["Lab. AH"]), kde=True, bins=30, ax=axes[1])
            axes[1].set_title("Log1p Distribution – Lab. AH")
            axes[1].set_xlabel("log1p(Lab. AH)")
        plt.tight_layout()
        artifacts.save_fig(plt, caption="Log Distributions – Eng. AH and Lab. AH")

    # Z vs robust-Z counts
    for col in targets:
        idx_z = z_outliers_idx(df[col], threshold=3.0)
        idx_rz = robust_z_outliers_idx(df[col], threshold=3.5)
        print(f"{col}: Z>3 => {len(idx_z)} outliers; robust Z>3.5 => {len(idx_rz)} outliers")

    # Outlier highlighting scatter
    if all(c in df.columns for c in ["Eng. AH", "Lab. AH"]):
        lo_e, hi_e = iqr_fences(df["Eng. AH"])
        lo_l, hi_l = iqr_fences(df["Lab. AH"])
        mask_e = (df["Eng. AH"] < lo_e) | (df["Eng. AH"] > hi_e)
        mask_l = (df["Lab. AH"] < lo_l) | (df["Lab. AH"] > hi_l)

        fig, ax = plt.subplots(figsize=(7, 5))
        sns.scatterplot(x="Eng. AH", y="Lab. AH", data=df, s=25, alpha=0.6, color="#9ecae1", label="Inliers", ax=ax)
        sns.scatterplot(x="Eng. AH", y="Lab. AH", data=df[mask_e | mask_l], s=40, color="#e34a33", label="Outliers (IQR)", ax=ax)
        ax.set_title("Eng vs Lab – Outliers highlighted (IQR)")
        ax.legend()
        artifacts.save_fig(plt, caption="Eng vs Lab – Outliers highlighted (IQR)")

    # Lab (>0) specific visuals & stats
    if "Lab. AH" in df.columns:
        lab_pos = df.loc[df["Lab. AH"] > 0, "Lab. AH"].dropna()
        if not lab_pos.empty:
            low_iqr, high_iqr = iqr_fences(lab_pos, k=1.5)
            x = np.log1p(lab_pos)
            med = np.median(x)
            mad = median_abs_deviation(x, scale="normal") or 1.0
            robust_z = 0.6745 * (x - med) / mad
            rz_idx = lab_pos.index[np.abs(robust_z) > 3.5]
            p99 = lab_pos.quantile(0.99)
            p995 = lab_pos.quantile(0.995)
            print({
                "Lab>0 n": lab_pos.size,
                "IQR_low": low_iqr, "IQR_high": high_iqr,
                "IQR_outliers(>high or <low)": int(((lab_pos < low_iqr) | (lab_pos > high_iqr)).sum()),
                "robustZ>3.5": len(rz_idx),
                "p99": p99, "p99.5": p995
            })

            fig, ax = plt.subplots(1, 2, figsize=(12, 4))
            sns.boxplot(y=lab_pos, ax=ax[0], color="#74c476")
            ax[0].set_title("Lab. AH (>0) – Boxplot")
            sns.histplot(np.log1p(lab_pos), kde=True, bins=30, ax=ax[1])
            ax[1].set_title("log1p(Lab. AH>0) – Distribution")
            plt.tight_layout()
            artifacts.save_fig(plt, caption="Lab. AH (>0) – Boxplot & log1p Distribution")

    # Eng log-robust-Z
    if "Eng. AH" in df.columns:
        eng = df["Eng. AH"].dropna()
        if not eng.empty:
            low_e, high_e = iqr_fences(eng, k=1.5)
            x = np.log(eng[eng > 0])
            med = np.median(x)
            mad = median_abs_deviation(x, scale="normal") or 1.0
            rz = 0.6745 * (x - med) / mad
            eng_rz_idx = x.index[np.abs(rz) > 3.5]
            print({
                "Eng IQR_high": high_e,
                "Eng IQR_outliers": int(((eng < low_e) | (eng > high_e)).sum()),
                "Eng robustZ>3.5 on log": len(eng_rz_idx)
            })

    # log(Eng) vs log1p(Lab)
    if all(c in df.columns for c in ["Eng. AH", "Lab. AH"]):
        xy = df[["Eng. AH", "Lab. AH"]].dropna().copy()
        if not xy.empty:
            xy["logEng"] = safe_log(xy["Eng. AH"])
            xy["logLab"] = np.log1p(xy["Lab. AH"])
            fig, ax = plt.subplots(figsize=(7, 5))
            sns.scatterplot(data=xy, x="logEng", y="logLab", s=15, alpha=0.5, ax=ax)
            ax.set_title("log(Eng) vs log1p(Lab)")
            artifacts.save_fig(plt, caption="log(Eng) vs log1p(Lab) scatter")


@capture_stdout
def step_categorical_plots(df: pd.DataFrame, artifacts: ReportArtifacts):
    # Bar plots for categorical variables (excluding targets)
    EXCLUDE_COLS = {"Eng. AH", "Lab. AH"}
    TOP_N = 15
    MAX_ALLOWED_UNIQUES = 200

    cat_vars = []
    for c in df.columns:
        if c in EXCLUDE_COLS:
            continue
        s = df[c]
        if not pd.api.types.is_numeric_dtype(s):
            cat_vars.append(c)
        else:
            if s.nunique(dropna=True) <= 20:
                cat_vars.append(c)

    for variable in cat_vars:
        s = df[variable].astype(str).fillna("Missing")
        nunique = s.nunique()
        if nunique > MAX_ALLOWED_UNIQUES:
            print(f"[SKIP] {variable}: {nunique} unique values (>{MAX_ALLOWED_UNIQUES}).")
            continue

        vc = s.value_counts()
        total = len(s)
        top = vc.head(TOP_N)
        other = vc.iloc[TOP_N:].sum()
        if other > 0:
            plot_counts = pd.concat([top, pd.Series({"Other": other})])
        else:
            plot_counts = top.copy()

        plot_pcts = (plot_counts / total) * 100
        fig, ax = plt.subplots(figsize=(10, 6))
        bars = ax.bar(plot_counts.index.astype(str), plot_counts.values, color="#4C78A8")
        for bar, count, pct in zip(bars, plot_counts.values, plot_pcts.values):
            h = bar.get_height()
            ax.text(bar.get_x() + bar.get_width() / 2, h, f"{int(count)} • {pct:.1f}%",
                    ha="center", va="bottom", fontsize=9)
        ax.set_title(f"Count Plot – {variable}")
        ax.set_xlabel(variable)
        ax.set_ylabel("Count")
        plt.xticks(rotation=30, ha="right")
        plt.tight_layout()
        artifacts.save_fig(plt, caption=f"Count Plot – {variable}")


@capture_stdout
def step_numeric_histograms(df: pd.DataFrame, artifacts: ReportArtifacts):
    EXCLUDE_COLS = {"Eng. AH", "Lab. AH"}
    num_vars = [c for c in df.columns
                if c not in EXCLUDE_COLS and pd.api.types.is_numeric_dtype(df[c])]
    num_vars = [c for c in num_vars if df[c].nunique(dropna=True) > 1]

    if len(num_vars) == 0:
        print("No numeric variables to plot.")
        return

    # Plot per-feature histograms laid out in a grid
    n = len(num_vars)
    ncols = 3
    nrows = int(np.ceil(n / ncols))
    fig, axes = plt.subplots(nrows=nrows, ncols=ncols, figsize=(5 * ncols, 3.5 * nrows))
    axes = axes.flatten() if n > 1 else [axes]

    for i, feature in enumerate(num_vars):
        ax = axes[i]
        sns.histplot(df[feature].dropna(), kde=True, bins=30, color="#5B8FF9", ax=ax)
        skew = df[feature].dropna().skew()
        ax.set_title(f"{feature}  |  Skewness: {skew:.2f}")
        ax.set_xlabel(feature)
        ax.set_ylabel("Count")

    # Hide any unused axes
    for j in range(i + 1, len(axes)):
        axes[j].axis("off")

    plt.tight_layout()
    artifacts.save_fig(plt, caption="Histograms for numeric variables")


@capture_stdout
def step_boolean_flags(df: pd.DataFrame, artifacts: ReportArtifacts):
    # Identify boolean-like flags
    EXCLUDE = {"Eng. AH", "Lab. AH", "Eng. SH", "Lab. SH"}
    candidate_cols = [c for c in df.columns if c not in EXCLUDE]

    bool_like_cols = []
    for c in candidate_cols:
        s = df[c]
        if pd.api.types.is_bool_dtype(s):
            bool_like_cols.append(c)
        elif pd.api.types.is_numeric_dtype(s):
            uniq = pd.unique(s.dropna())
            if set(uniq).issubset({0, 1}):
                bool_like_cols.append(c)

    if len(bool_like_cols) == 0:
        print("No boolean/indicator columns found.")
        return

    # True-rate for each flag
    rows = []
    for c in bool_like_cols:
        s = df[c].astype(float)
        n = s.notna().sum()
        true_rate = (s.fillna(0).mean()) * 100.0
        rows.append({"flag": c, "true_rate_pct": true_rate, "non_missing": int(n)})

    true_rates = pd.DataFrame(rows).sort_values("true_rate_pct", ascending=False).reset_index(drop=True)
    artifacts.add_table("Flag True-rates (%)", true_rates)

    # Plot Top-15 True-rates
    top_n = 15
    top15 = true_rates.head(top_n).copy()
    fig, ax = plt.subplots(figsize=(9, max(5, 0.4 * len(top15))))
    sns.barplot(data=top15, y="flag", x="true_rate_pct", color="#4C78A8", orient="h", ax=ax)
    for i, (rate, nm) in enumerate(zip(top15["true_rate_pct"], top15["non_missing"])):
        ax.text(rate + 0.5, i, f"{rate:.1f}%  (n={nm})", va="center", fontsize=9)
    ax.set_title("Top 15 Flags by True-rate (%)")
    ax.set_xlabel("True-rate (%)")
    ax.set_ylabel("Flag")
    ax.set_xlim(0, max(100, top15["true_rate_pct"].max() + 5))
    plt.tight_layout()
    artifacts.save_fig(plt, caption="Top 15 Flags by True-rate (%)")

    # Count of unbalanced flags (True-rate < 5%)
    unbalanced_threshold = 5.0
    n_unbalanced = int((true_rates["true_rate_pct"] < unbalanced_threshold).sum())
    n_total_flags = len(true_rates)
    print(f"Flags with True-rate < {unbalanced_threshold:.1f}%: {n_unbalanced} out of {n_total_flags} flags")


@capture_stdout
def step_multivariate(df: pd.DataFrame, artifacts: ReportArtifacts):
    print("\n##########################################")
    print("########## Multivariate Analysis #########")
    print("##########################################\n")

    # Average Eng. AH and Lab. AH across non-dominant categories
    if all(c in df.columns for c in ["Eng. AH", "Lab. AH"]):
        print("Average Eng. AH and Lab. AH by categorical features")
        exclude = {"Eng. AH", "Lab. AH"}
        for col in [c for c in df.columns if c not in exclude]:
            s = df[col].copy()
            s_cat = s.astype(object).where(~s.isna(), "Missing")
            counts = s_cat.value_counts(dropna=False)
            pct = (counts / len(s_cat)) * 100
            if pct.max() > 95:
                continue  # skip dominant columns
            df_aux = pd.DataFrame({
                col: s_cat,
                "Eng. AH": df["Eng. AH"],
                "Lab. AH": df["Lab. AH"]
            })
            g = (df_aux.groupby(col, dropna=False).agg(
                **{
                    "% of Rows": (col, lambda x: 100 * len(x) / len(df_aux)),
                    "Mean Eng. AH": ("Eng. AH", "mean"),
                    "Mean Lab. AH": ("Lab. AH", "mean")
                }
            ).reset_index().rename(columns={col: "Value"}).sort_values("% of Rows", ascending=False))
            g["% of Rows"] = g["% of Rows"].round(1)
            g["Mean Eng. AH"] = g["Mean Eng. AH"].round(3)
            g["Mean Lab. AH"] = g["Mean Lab. AH"].round(3)
            print(f"\n=== {col} ===")
            print(g.to_string(index=False))

    # Kruskal–Wallis for some key categoricals (if present)
    cats = [c for c in ["CCN_Data Hub", "type_of_investigation", "Investigation_type", "Region"] if c in df.columns]
    for col in cats:
        s = df[col].astype(str).fillna("Missing")
        if "Eng. AH" in df.columns:
            y_eng = safe_log(df["Eng. AH"])
            groups_eng = [y_eng[s == level].dropna() for level in s.unique()]
            H_eng, p_eng = sstats.kruskal(*groups_eng)
            n_eng = int(sum(map(len, groups_eng)))
            k_eng = len(groups_eng)
            eps2_eng = (H_eng - (k_eng - 1)) / (n_eng - 1) if n_eng > 1 else np.nan
        else:
            H_eng = p_eng = eps2_eng = np.nan
            n_eng = k_eng = 0

        if "Lab. AH" in df.columns:
            y_lab_pos = df.loc[df["Lab. AH"] > 0, "Lab. AH"]
            y_lab_pos_log1p = np.log1p(y_lab_pos)
            s_pos = s.loc[y_lab_pos.index]
            groups_lab = [y_lab_pos_log1p[s_pos == level].dropna() for level in s_pos.unique()]
            if sum(map(len, groups_lab)) > 0 and len(groups_lab) > 1:
                H_lab, p_lab = sstats.kruskal(*groups_lab)
                n_lab = int(sum(map(len, groups_lab)))
                k_lab = len(groups_lab)
                eps2_lab = (H_lab - (k_lab - 1)) / (n_lab - 1) if n_lab > 1 else np.nan
            else:
                H_lab = p_lab = eps2_lab = np.nan
        else:
            H_lab = p_lab = eps2_lab = np.nan

        print(f"\n[{col}]  Kruskal–Wallis:")
        print(f"Eng.AH: H={H_eng:.2f}, p={p_eng:.3e}, epsilon^2={eps2_eng:.3f}  (n={n_eng}, k={k_eng})")
        print(f"Lab.AH(>0): H={H_lab:.2f}, p={p_lab:.3e}, epsilon^2={eps2_lab:.3f}")

    # Mann–Whitney for boolean flags
    flags = []
    for c in df.columns:
        if c in {"Eng. AH", "Lab. AH", "Eng. SH", "Lab. SH"}:
            continue
        s = df[c]
        if pd.api.types.is_bool_dtype(s) or (pd.api.types.is_numeric_dtype(s) and set(pd.unique(s.dropna())).issubset({0, 1})):
            if 0 < (s == 1).mean() < 1:
                flags.append(c)

    rows = []
    for f in flags:
        s = df[f]
        # Eng AH effect
        if "Eng. AH" in df.columns:
            x = df.loc[s == 1, "Eng. AH"].dropna()
            y = df.loc[s == 0, "Eng. AH"].dropna()
            if len(x) > 0 and len(y) > 0:
                U, p = sstats.mannwhitneyu(x, y, alternative="two-sided")
                r_rb = 1 - (2 * U) / (len(x) * len(y))
            else:
                p = r_rb = np.nan
        else:
            p = r_rb = np.nan

        # Lab AH (>0) on log1p
        if "Lab. AH" in df.columns:
            xl = df.loc[(s == 1) & (df["Lab. AH"] > 0), "Lab. AH"].dropna()
            yl = df.loc[(s == 0) & (df["Lab. AH"] > 0), "Lab. AH"].dropna()
            if len(xl) > 0 and len(yl) > 0:
                U2, p2 = sstats.mannwhitneyu(np.log1p(xl), np.log1p(yl), alternative="two-sided")
                r_rb2 = 1 - (2 * U2) / (len(xl) * len(yl))
            else:
                p2 = r_rb2 = np.nan
        else:
            p2 = r_rb2 = np.nan

        rows.append({"flag": f, "p_Eng": p, "r_rb_Eng": r_rb, "p_Lab_pos": p2, "r_rb_Lab_pos": r_rb2,
                     "True_rate_%": 100 * (s == 1).mean()})

    if rows:
        effects = pd.DataFrame(rows).sort_values("r_rb_Eng", key=lambda s: s.abs(), ascending=False)
        print("\nTop flags by |effect on Eng. AH|:")
        print(effects.head(15).to_string(index=False))
        # Add counts & FDR filtering as in original
        effects["n_true"] = effects["True_rate_%"] * len(df) / 100.0
        effects["n_false"] = len(df) - effects["n_true"]
        min_true = 20
        min_rate = 1.0
        mask_support = (effects["n_true"] >= min_true) & (effects["True_rate_%"] >= min_rate)
        eff_filt = effects.loc[mask_support].copy()
        if not eff_filt.empty and eff_filt["p_Eng"].notna().any():
            pvals = eff_filt["p_Eng"].fillna(1.0).values
            m = len(pvals)
            order = np.argsort(pvals)
            rank = np.empty_like(order)
            rank[order] = np.arange(1, m + 1)
            q = pvals * m / rank
            # Monotone correction
            q_sorted = q[np.argsort(order)]
            for i in range(len(q_sorted) - 2, -1, -1):
                q_sorted[i] = min(q_sorted[i], q_sorted[i + 1])
            eff_filt["q_Eng_FDR"] = q_sorted
            eff_final = (eff_filt.sort_values(
                ["r_rb_Eng", "q_Eng_FDR"],
                key=lambda s: s.abs() if s.name == "r_rb_Eng" else s,
                ascending=[False, True]
            ))
            print("\nFlags (well-supported) ranked by |effect on Eng. AH| with FDR control:")
            cols = ["flag", "True_rate_%", "n_true", "r_rb_Eng", "p_Eng", "q_Eng_FDR", "r_rb_Lab_pos", "p_Lab_pos"]
            print(eff_final[cols].head(15).to_string(index=False))
            artifacts.add_table("Flag Effects (FDR on Eng)", eff_final[cols])
        else:
            artifacts.add_table("Flag Effects (raw)", effects)

    # Heatmaps for median Eng. AH by type_of_investigation × Region
    if all(c in df.columns for c in ["type_of_investigation", "Region", "Eng. AH"]):
        pivot = pd.pivot_table(df, index="type_of_investigation", columns="Region", values="Eng. AH", aggfunc="median")
        fig, ax = plt.subplots(figsize=(7, 4))
        sns.heatmap(pivot, annot=True, fmt=".1f", cmap="YlGnBu", ax=ax)
        ax.set_title("Median Eng. AH by type_of_investigation × Region")
        plt.tight_layout()
        artifacts.save_fig(plt, caption="Median Eng. AH by type_of_investigation × Region")

    # Lab>0 rate heatmap
    if all(c in df.columns for c in ["type_of_investigation", "Region", "Lab. AH"]):
        rate = pd.pivot_table(
            df.assign(lab_has=(df["Lab. AH"] > 0).astype(int)),
            index="type_of_investigation", columns="Region", values="lab_has", aggfunc="mean"
        ) * 100
        fig, ax = plt.subplots(figsize=(7, 4))
        sns.heatmap(rate, annot=True, fmt=".1f", cmap="PuRd", ax=ax)
        ax.set_title("% with Lab>0 by type_of_investigation × Region")
        plt.tight_layout()
        artifacts.save_fig(plt, caption="% with Lab>0 by type_of_investigation × Region")

    # Numeric correlates (Spearman) with logged targets
    num_feats = [c for c in df.columns
                 if c not in {"Eng. AH", "Lab. AH", "Eng. SH", "Lab. SH"}
                 and pd.api.types.is_numeric_dtype(df[c])
                 and df[c].nunique(dropna=True) > 1]
    rows = []
    for c in num_feats:
        s = df[c].astype(float)
        r1 = sstats.spearmanr(s, safe_log(df["Eng. AH"]) if "Eng. AH" in df.columns else s * np.nan, nan_policy="omit").correlation
        r2 = sstats.spearmanr(s, np.log1p(df["Lab. AH"].fillna(0)) if "Lab. AH" in df.columns else s * np.nan, nan_policy="omit").correlation
        rows.append({"feature": c, "spearman_log_eng": r1, "spearman_log1p_lab": r2})
    if rows:
        corr_num = pd.DataFrame(rows)
        print("\nTop numeric correlates with log(Eng. AH):")
        print(corr_num.sort_values("spearman_log_eng", key=np.abs, ascending=False).head(10))
        print("\nTop numeric correlates with log1p(Lab. AH):")
        print(corr_num.sort_values("spearman_log1p_lab", key=np.abs, ascending=False).head(10))
        artifacts.add_table("Numeric Spearman Correlates", corr_num.sort_values("spearman_log_eng", key=np.abs, ascending=False).head(20))

        # Regressions plots for top Eng correlates
        if "Eng. AH" in df.columns:
            topE = corr_num.sort_values("spearman_log_eng", key=np.abs, ascending=False).head(6)["feature"].tolist()
            fig, axes = plt.subplots(2, 3, figsize=(15, 8))
            axes = axes.flatten()
            for i, f in enumerate(topE):
                sns.regplot(x=df[f], y=safe_log(df["Eng. AH"]), lowess=True,
                            scatter_kws=dict(s=12, alpha=0.4), line_kws=dict(color="red"), ax=axes[i])
                axes[i].set_xlabel(f); axes[i].set_ylabel("log(Eng. AH)")
            plt.tight_layout()
            artifacts.save_fig(plt, caption="Top numeric correlates vs log(Eng. AH)")

    # Joint effects: type_of_investigation × total_test_count (bins)
    if "total_test_count" in df.columns and "type_of_investigation" in df.columns:
        tt = df["total_test_count"]
        tt_bins = pd.cut(tt, bins=[-0.5, 0, 2, 5, 10, 26],
                         labels=["0", "1-2", "3-5", "6-10", "11-26"])
        cat = df["type_of_investigation"].astype(str).fillna("Missing")

        if "Eng. AH" in df.columns:
            piv_eng = pd.pivot_table(df, index=cat, columns=tt_bins, values="Eng. AH", aggfunc="median").fillna(0)
            fig, ax = plt.subplots(figsize=(8, 5))
            sns.heatmap(piv_eng, annot=True, fmt=".1f", cmap="YlGnBu", ax=ax)
            ax.set_title("Median Eng. AH by type_of_investigation × total_test_count (bins)")
            ax.set_xlabel("total_test_count (bins)")
            ax.set_ylabel("type_of_investigation")
            plt.tight_layout()
            artifacts.save_fig(plt, caption="Median Eng. AH by type_of_investigation × total_test_count (bins)")

        if "Lab. AH" in df.columns:
            lab_has = (df["Lab. AH"] > 0).astype(int)
            piv_lab_rate = pd.pivot_table(
                pd.DataFrame({"cat": cat, "bin": tt_bins, "lab_has": lab_has}),
                index="cat", columns="bin", values="lab_has", aggfunc="mean"
            ) * 100
            fig, ax = plt.subplots(figsize=(8, 5))
            sns.heatmap(piv_lab_rate, annot=True, fmt=".1f", cmap="PuRd", ax=ax)
            ax.set_title("% with Lab>0 by type_of_investigation × total_test_count (bins)")
            ax.set_xlabel("total_test_count (bins)")
            ax.set_ylabel("type_of_investigation")
            plt.tight_layout()
            artifacts.save_fig(plt, caption="% with Lab>0 by type_of_investigation × total_test_count (bins)")

        # Facet by Region (heatmap with multi-columns)
        if "Region" in df.columns and "Eng. AH" in df.columns:
            reg = df["Region"].astype(str).fillna("Missing")
            gdf = pd.DataFrame({"Eng": df["Eng. AH"], "cat": cat, "bin": tt_bins, "Region": reg})
            g = gdf.groupby(["Region", "cat", "bin"], dropna=False)["Eng"].median().reset_index()
            g = g.pivot_table(index=["cat"], columns=["bin", "Region"], values="Eng").fillna(0)
            fig, ax = plt.subplots(figsize=(12, 6))
            sns.heatmap(g, annot=False, cmap="YlGnBu", ax=ax)
            ax.set_title("Median Eng. AH by type_of_investigation × total_test_count (bins) × Region")
            plt.tight_layout()
            artifacts.save_fig(plt, caption="Median Eng. AH by type_of_investigation × bins × Region")

    # Cramér’s V among categoricals/flags
    cat_cols = [c for c in ["type_of_investigation", "Investigation_type", "Region", "CCN_Data Hub"] if c in df.columns]
    flag_candidates = [
        "marking_test", "abnormal_operations_test", "input_test",
        "heating_test", "impact_test", "steady_force_test",
        "connector_overload_test"
    ]
    cat_cols += [c for c in flag_candidates if c in df.columns]
    if len(cat_cols) >= 2:
        C = pd.DataFrame(index=cat_cols, columns=cat_cols, dtype=float)
        for i, ci in enumerate(cat_cols):
            for j, cj in enumerate(cat_cols):
                if i == j:
                    C.iloc[i, j] = 1.0
                else:
                    C.iloc[i, j] = cramers_v(
                        df[ci].astype(str).fillna("Missing"),
                        df[cj].astype(str).fillna("Missing")
                    )
        fig, ax = plt.subplots(figsize=(8, 6))
        sns.heatmap(C, vmin=0, vmax=1, cmap="coolwarm", ax=ax)
        ax.set_title("Cramér’s V among categorical variables/flags")
        plt.tight_layout()
        artifacts.save_fig(plt, caption="Cramér’s V heatmap (categoricals/flags)")

        # Cluster map
        sns.clustermap(C, vmin=0, vmax=1, cmap="coolwarm", figsize=(9, 9))
        artifacts.save_fig(plt, caption="Cramér’s V clustermap (categoricals/flags)")

    # Numeric correlation (Spearman)
    base_num = [c for c in ["Eng. AH", "Lab. AH", "total_test_count", "total_CB_count", "standard_count"] if c in df.columns]
    num_cols = base_num + [
        c for c in df.columns
        if pd.api.types.is_numeric_dtype(df[c])
        and c not in base_num
        and c not in ["Eng. AH", "Lab. AH"]
        and df[c].nunique(dropna=True) > 1
        and df[c].nunique(dropna=True) < 100
    ]
    if num_cols:
        corr = df[num_cols].corr(method="spearman")
        fig, ax = plt.subplots(figsize=(10, 8))
        sns.heatmap(corr, cmap="RdBu_r", vmin=-1, vmax=1, center=0, ax=ax)
        ax.set_title("Spearman correlation matrix – numeric variables")
        plt.tight_layout()
        artifacts.save_fig(plt, caption="Spearman correlation matrix – numeric variables")

        sns.clustermap(corr, cmap="RdBu_r", center=0, figsize=(10, 10))
        artifacts.save_fig(plt, caption="Spearman clustermap – numeric variables")

    # Conditional effect of flags within total_test_count bins
    flags = [f for f in ["marking_test", "abnormal_operations_test", "input_test", "heating_test", "impact_test"] if f in df.columns]
    if flags and "total_test_count" in df.columns:
        tt_bins = pd.cut(df["total_test_count"], bins=[-0.5, 0, 2, 5, 10, 26], labels=["0", "1-2", "3-5", "6-10", "11-26"])
        for f in flags:
            cols_present = ["Eng. AH", "Lab. AH"]
            if not any(c in df.columns for c in cols_present):
                continue
            dfp = pd.DataFrame({
                "bin": tt_bins,
                "flag": df[f].astype("float").fillna(0).astype(int),
                "Eng": df["Eng. AH"] if "Eng. AH" in df.columns else np.nan,
                "Lab": df["Lab. AH"] if "Lab. AH" in df.columns else np.nan
            }).dropna(subset=["bin"])
            med = dfp.groupby(["bin", "flag"]).agg(med_Eng=("Eng", "median"), med_Lab=("Lab", "median")).reset_index()
            pe = med.pivot_table(index="bin", columns="flag", values="med_Eng")
            pl = med.pivot_table(index="bin", columns="flag", values="med_Lab")
            pe["delta_Eng"] = pe.get(1, np.nan) - pe.get(0, np.nan)
            pl["delta_Lab"] = pl.get(1, np.nan) - pl.get(0, np.nan)

            fig, axes = plt.subplots(1, 2, figsize=(10, 3.5))
            axes[0].plot(pe.index.astype(str), pe["delta_Eng"], marker="o", color="#E15759")
            axes[0].axhline(0, color="grey", lw=1)
            axes[0].set_title(f"{f}: Median Δ Eng (True−False) by bin")
            axes[0].set_xlabel("total_test_count (bins)")
            axes[0].set_ylabel("Median Δ Eng. AH")

            axes[1].plot(pl.index.astype(str), pl["delta_Lab"], marker="s", color="#76B7B2")
            axes[1].axhline(0, color="grey", lw=1)
            axes[1].set_title(f"{f}: Median Δ Lab (True−False) by bin")
            axes[1].set_xlabel("total_test_count (bins)")
            axes[1].set_ylabel("Median Δ Lab. AH")

            plt.tight_layout()
            artifacts.save_fig(plt, caption=f"{f}: Conditional median deltas by total_test_count bins")

    # Mutual Information (requires imputation/encoding)
    feat_pool = []
    feat_pool += [c for c in ["total_test_count", "total_CB_count", "standard_count"] if c in df.columns]
    feat_pool += [c for c in ["marking_test", "abnormal_operations_test", "input_test",
                              "heating_test", "impact_test", "steady_force_test",
                              "connector_overload_test"] if c in df.columns]
    feat_pool += [c for c in ["Region", "type_of_investigation", "Investigation_type"] if c in df.columns]
    if feat_pool:
        X_parts = []
        for c in feat_pool:
            s = df[c]
            if pd.api.types.is_numeric_dtype(s):
                arr = s.to_numpy().reshape(-1, 1)
                arr = np.where(np.isnan(arr), np.nanmedian(arr), arr)
                X_parts.append(arr)
            else:
                oh = pd.get_dummies(s.astype(str).fillna("Missing"), drop_first=False)
                arr = oh.to_numpy()
                X_parts.append(arr)
        X = np.hstack(X_parts)

        if "Eng. AH" in df.columns:
            y_eng = df["Eng. AH"].to_numpy()
            mi_eng = mutual_info_regression(X, y_eng, random_state=RANDOM_STATE)
            idx_top_eng = np.argsort(mi_eng)[::-1][:20]
            print("\nMutual information (Eng. AH) – top 20 dimensions of X (values only):")
            print(mi_eng[idx_top_eng])

        if "Lab. AH" in df.columns:
            y_lab_has = (df["Lab. AH"] > 0).astype(int).to_numpy()
            mi_lab = mutual_info_classif(X, y_lab_has, random_state=RANDOM_STATE)
            idx_top_lab = np.argsort(mi_lab)[::-1][:20]
            print("\nMutual information (Lab>0) – top 20 dimensions of X (values only):")
            print(mi_lab[idx_top_lab])

    # Partial correlations (residualization by total_test_count)
    if "total_test_count" in df.columns:
        def _residualize(v, control):
            v = pd.Series(v).astype(float)
            c = pd.Series(control).astype(float)
            Xc = np.column_stack([np.ones_like(c), c])
            v_filled = v.fillna(v.median()).values
            beta, _, _, _ = np.linalg.lstsq(Xc, v_filled, rcond=None)
            return v - Xc.dot(beta)

        control = df["total_test_count"].fillna(0)
        cand_num = [c for c in ["Eng. AH", "Lab. AH", "total_CB_count", "standard_count"] if c in df.columns]
        if cand_num:
            res_mat = {c: _residualize(df[c], control) for c in cand_num}
            cols = list(res_mat.keys())
            R = pd.DataFrame(index=cols, columns=cols, dtype=float)
            for i, ci in enumerate(cols):
                for j, cj in enumerate(cols):
                    R.iloc[i, j] = sstats.spearmanr(res_mat[ci], res_mat[cj], nan_policy="omit").correlation
            fig, ax = plt.subplots(figsize=(6, 5))
            sns.heatmap(R, cmap="RdBu_r", center=0, vmin=-1, vmax=1, annot=True, fmt=".2f", ax=ax)
            ax.set_title("Spearman among residuals (control: total_test_count)")
            plt.tight_layout()
            artifacts.save_fig(plt, caption="Spearman among residuals | control = total_test_count")

        # Example partial correlation for standard_count vs Eng. AH controlling for total_test_count
        if "standard_count" in df.columns and "Eng. AH" in df.columns:
            x = df["total_test_count"].fillna(0).values
            Xc = np.column_stack([np.ones_like(x), x])  # intercept + control
            def residualize_vec(y):
                yv = pd.Series(y).astype(float).fillna(pd.Series(y).median()).values
                beta, _, _, _ = np.linalg.lstsq(Xc, yv, rcond=None)
                return yv - Xc.dot(beta)
            res_t = residualize_vec(safe_log(df["Eng. AH"]))
            res_f = residualize_vec(df["standard_count"])
            partial_r = np.corrcoef(res_t, res_f)[0, 1]
            print("Partial correlation (log Eng. AH ~ standard_count | total_test_count):", round(partial_r, 3))


# ==================================================
# 4. Main
# ==================================================
def parse_args():
    parser = argparse.ArgumentParser(description="EDA with full Word report export.")
    parser.add_argument("--input", type=str, default="1. Data.xlsx", help="Path to input Excel file")
    parser.add_argument("--output", type=str, default="EDA_Report.docx", help="Path to output Word (.docx)")
    parser.add_argument("--report_dir", type=str, default="reports", help="Directory to store figures and assets")
    return parser.parse_args()


def main():
    args = parse_args()
    input_path = Path(args.input).resolve()
    output_docx = Path(args.output).resolve()
    report_dir = Path(args.report_dir).resolve()
    report_dir.mkdir(parents=True, exist_ok=True)

    # Artifacts tracker
    artifacts = ReportArtifacts(report_dir)

    # Read data
    df = read_data(input_path)

    # Run EDA steps
    step_preliminary_info(df, artifacts=artifacts)
    step_univariate(df, artifacts=artifacts)
    step_categorical_plots(df, artifacts=artifacts)
    step_numeric_histograms(df, artifacts=artifacts)
    step_boolean_flags(df, artifacts=artifacts)
    step_multivariate(df, artifacts=artifacts)

    # Write Word report
    metadata = {
        "Author": "Nuno Monteiro and Celine Androun",
        "Targets": "Eng. AH, Lab. AH",
        "Input File": str(input_path.name)
    }
    write_word_report(
        artifacts=artifacts,
        output_docx=output_docx,
        title="Exploratory Data Analysis Report – Eng. AH & Lab. AH",
        meta=metadata
    )

    print(f"\nReport successfully written to: {output_docx}")
    print(f"Figures saved under: {artifacts.fig_dir}")


if __name__ == "__main__":
    main()
